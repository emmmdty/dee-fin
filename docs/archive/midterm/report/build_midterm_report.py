#!/usr/bin/env python
"""Generate the Fin-EKG midterm report from the official DOCX template."""
# ruff: noqa: E501

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = Path(__file__).resolve().parent
MIDTERM = HERE.parent
DATA = MIDTERM / "data"
FIGURES = MIDTERM / "figures"
TEMPLATE = HERE / "童佳锴中期报告初稿.docx"  # 底板：封面已填好，标题由脚本替换
# 直接产出完整报告（封面+正文），不再需要组装脚本
OUT_DOCX = HERE / "童佳锴中期报告初稿_修改版.docx"

NEW_TITLE = "基于约束引导的事件知识图谱构建与时序推理方法研究"
OLD_TITLE = "面向金融事件图谱的可验证构建与时序推理"

TABLE_CAPTIONS = [
    "表1 三部分研究内容中期进展总表",
    "表2 数据集与评测口径清单",
    "表3 SARGE 与公开基线对比",
    "表4 SARGE 自建中文事件节点统计",
    "表5 监督微调与 GRPO-RLVR 关系抽取结果",
    "表6 ICEWS14 时序预测结果",
    "表7 ICEWS14 与近年方法口径对照",
    "表8 保形预测风险控制结果",
    "表9 后续研究进度安排",
]

FIGURE_CAPTIONS = [
    "图1 可验证的事件知识图谱构建与时序推理总体框架",
    "图2 本文方法与公开基线在中文文档级事件抽取上的性能对比",
    "图3 自建中文事件图谱的事件类型分布",
    "图4 监督微调与 GRPO-RLVR 的事件共指质量对比",
    "图5 自建中文事件图谱整体概览",
    "图6 中文事件图谱的构建与一致性闭包过程",
    "图7 新奥股份局部事件中心知识图谱实例",
    "图8 ICEWS14 上各方法的时序预测性能",
    "图9 Path-RL 的训练过程",
    "图10 不同校准策略下的覆盖率与漂移失真",
]


_CJK = "一-鿿"


def _squeeze(text: str) -> str:
    """中英/中数之间不留空格(GB 排版)：删除中文与 ASCII 字母/数字/% 相邻的空格。
    标题与图题里"图N /表N /N.N "后的那个空格由 _squeeze_title 单独保留。"""
    prev = None
    while prev != text:
        prev = text
        text = re.sub(rf"([{_CJK}])[ \t]+([0-9A-Za-z%])", r"\1\2", text)
        text = re.sub(rf"([0-9A-Za-z%）)])[ \t]+([{_CJK}])", r"\1\2", text)
    return text


def _squeeze_title(text: str) -> str:
    """标题/图题：保留"图N /表N /N.N "这一处数字后的空格，其余按 _squeeze 收紧。"""
    match = re.match(r"^((?:图|表)\s*\d+|\d+(?:\.\d+)*)\s+", text)
    if match:
        return match.group(0) + _squeeze(text[match.end():])
    return _squeeze(text)


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def _round4(value: float) -> float:
    return round(float(value), 4)


def _fc_metrics(data: dict) -> dict:
    return data["metrics"] if "metrics" in data else data


def parse_fusion_sweep() -> dict[str, float | str]:
    text = (DATA / "fuse_re_gcn.log").read_text(encoding="utf-8")
    match = re.search(r"BEST filtered MRR=([0-9.]+) at ws=([0-9.]+) wc=([0-9.]+)", text)
    if not match:
        raise ValueError("Cannot parse fusion sweep best row from fuse_re_gcn.log")
    best_mrr, ws, wc = match.groups()
    row = re.search(
        rf"ws={re.escape(ws)} wc={re.escape(wc)}\s+MRR=([0-9.]+)\s+H@1=([0-9.]+)\s+H@10=([0-9.]+)",
        text,
    )
    if not row:
        raise ValueError("Cannot parse fusion sweep metric row from fuse_re_gcn.log")
    _, h1, h10 = row.groups()
    return {
        "mrr_tfilt": _round4(float(best_mrr)),
        "hits@1_tfilt": _round4(float(h1)),
        "hits@10_tfilt": _round4(float(h10)),
        "weights": f"ws={ws}, wc={wc}",
    }


def load_ch3_results() -> dict[str, dict[str, float | str]]:
    freq = _fc_metrics(_load_json(DATA / "freq_icews14.json"))
    tgnn = _fc_metrics(_load_json(DATA / "tgnn_icews14" / "metrics.json"))
    path = _fc_metrics(_load_json(DATA / "path_rl_icews14" / "eval.json"))
    regcn = _load_json(DATA / "re_gcn_icews14" / "metrics.json")
    hybrid = _load_json(DATA / "hybrid_icews14" / "metrics.json")
    return {
        "frequency": {
            "mrr": _round4(freq["mrr"]),
            "hits@1": _round4(freq["hits@1"]),
            "hits@10": _round4(freq["hits@10"]),
        },
        "temporal_gnn": {
            "mrr": _round4(tgnn["mrr"]),
            "hits@1": _round4(tgnn["hits@1"]),
            "hits@10": _round4(tgnn["hits@10"]),
        },
        # recurrency 为 CPU 多信号 copy 基线，无独立 metrics.json 落盘；
        # 数值来源 docs/midterm/RESULTS_SUMMARY_2026-06.md（与正文、图保持一致）。
        "recurrency": {"mrr": 0.3560, "hits@1": 0.2830, "hits@10": 0.4880},
        "path_rl": {
            "mrr": _round4(path["mrr"]),
            "hits@1": _round4(path["hits@1"]),
            "hits@10": _round4(path["hits@10"]),
        },
        "re_gcn": {
            "mrr_tfilt": _round4(regcn["mrr_tfilt"]),
            "hits@1_tfilt": _round4(regcn["hits@1_tfilt"]),
            "hits@10_tfilt": _round4(regcn["hits@10_tfilt"]),
        },
        "hybrid_single": {
            "mrr_tfilt": _round4(hybrid["mrr_tfilt"]),
            "hits@1_tfilt": _round4(hybrid["hits@1_tfilt"]),
            "hits@10_tfilt": _round4(hybrid["hits@10_tfilt"]),
        },
        "fusion_sweep": parse_fusion_sweep(),
    }


def load_conformal_results() -> dict[str, dict[str, float]]:
    data = _load_json(DATA / "conformal_gnn_icews14.json")
    return {
        name: {
            "coverage": _round4(metrics["conformal_coverage"]),
            "drift_gap": _round4(metrics["coverage_drift_gap"]),
            "set_size": round(float(metrics["conformal_set_size"]), 1),
        }
        for name, metrics in data["calibrators"].items()
    }


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def add_text(
    paragraph,
    text: str,
    *,
    font: str = "宋体",
    size: float = 12,
    bold: bool = False,
    color: str | None = None,
):
    run = paragraph.add_run(text)
    set_east_asia_font(run, font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def set_paragraph_body(paragraph, *, first_line: bool = True) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if first_line:
        fmt.first_line_indent = Pt(24)
    else:
        fmt.first_line_indent = Pt(0)


def add_body_paragraph(doc: Document, text: str, *, first_line: bool = True):
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=first_line)
    add_text(p, _squeeze(text), size=12)
    return p


def add_note_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=False)
    p.paragraph_format.left_indent = Pt(18)
    add_text(p, _squeeze(text), size=10.5, color="555555")
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.first_line_indent = None
    size = {1: 15, 2: 14, 3: 12}.get(level, 12)
    add_text(p, _squeeze_title(text), font="黑体", size=size, bold=True, color="000000")
    return p


def add_center_title(doc: Document, text: str, *, size: int = 16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(12)
    add_text(p, text, font="黑体", size=size, bold=True, color="000000")
    return p


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def set_cell_edges(cell, edges: dict[str, int | None]) -> None:
    """按 GB/T 7714 三线表需要，仅画指定的单元格边线。

    edges 映射边名(top/bottom/left/right/insideH/insideV) -> 线宽(单位 1/8 磅)，
    值为 None 表示隐藏该边。
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in list(tc_pr.findall(qn("w:tcBorders"))):
        tc_pr.remove(old)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        size = edges.get(edge)
        if size is not None:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), str(size))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")
        borders.append(element)
    tc_pr.insert(0, borders)


def apply_three_line_table(table) -> None:
    """把表格渲染成三线表：顶线、表头底线、底线（约 1.5/1.0/1.5 磅），无竖线与其余横线。"""
    outer, header = 12, 8  # 1/8 磅：12≈1.5 磅，8≈1.0 磅
    rows = table.rows
    n = len(rows)
    for ri, row in enumerate(rows):
        for cell in row.cells:
            edges: dict[str, int | None] = {
                "left": None,
                "right": None,
                "insideH": None,
                "insideV": None,
                "top": None,
                "bottom": None,
            }
            if ri == 0:
                edges["top"] = outer
                edges["bottom"] = outer if n == 1 else header
            if ri == n - 1 and ri != 0:
                edges["bottom"] = outer
            set_cell_edges(cell, edges)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    width = OxmlElement("w:tcW")
    width.set(qn("w:w"), str(int(width_cm * 567)))
    width.set(qn("w:type"), "dxa")
    tc_pr.append(width)


def set_cell_text(cell, text: str, *, bold: bool = False, width_cm: float | None = None) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if width_cm is not None:
        set_cell_width(cell, width_cm)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph(paragraph)
    add_text(paragraph, _squeeze(text), size=10.0, bold=bold)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    caption: str,
    *,
    widths_cm: list[float] | None = None,
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_keep_with_next(p)
    add_text(p, _squeeze_title(caption), font="宋体", size=10.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_row_cant_split(table.rows[0])
    for i, header in enumerate(headers):
        set_cell_text(
            table.rows[0].cells[i],
            header,
            bold=True,
            width_cm=widths_cm[i] if widths_cm else None,
        )
    for row in rows:
        cells = table.add_row().cells
        set_row_cant_split(table.rows[-1])
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, width_cm=widths_cm[i] if widths_cm else None)
    apply_three_line_table(table)
    doc.add_paragraph()


def set_keep_with_next(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_figure(doc: Document, image: Path, caption: str, width_cm: float = 13.5) -> None:
    if not image.exists():
        add_body_paragraph(doc, f"（图像文件缺失：{image.name}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap, _squeeze_title(caption), font="宋体", size=10.5)


def replace_title_and_clear_body(doc: Document) -> None:
    """替换封面标题，然后清除旧正文（从第一个 Heading1 '1 ' 起），保留封面和说明页。"""
    # 1. 替换标题
    for p in doc.paragraphs:
        if OLD_TITLE in p.text:
            for run in p.runs:
                if OLD_TITLE in run.text:
                    run.text = run.text.replace(OLD_TITLE, NEW_TITLE)
                    break
            break

    # 2. 找到第一个正文 Heading1（以 "1 " 开头）
    body = doc.element.body
    first_h1 = None
    for child in list(body):
        if child.tag == qn("w:p"):
            style_el = child.find(qn("w:pPr"))
            if style_el is not None:
                pstyle = style_el.find(qn("w:pStyle"))
                if pstyle is not None and pstyle.get(qn("w:val")) == "1":
                    texts = []
                    for rt in child.iter(qn("w:t")):
                        if rt.text:
                            texts.append(rt.text)
                    full = "".join(texts).strip()
                    if full.startswith("1 ") or full.startswith("1\t"):
                        first_h1 = child
                        break

    # 3. 删除 first_h1 及之后的所有元素（除 sectPr）
    if first_h1 is not None:
        removing = False
        for child in list(body):
            if child is first_h1:
                removing = True
            if removing and child.tag != qn("w:sectPr"):
                body.remove(child)





def add_progress_section(doc: Document, ch3: dict[str, dict[str, float | str]], conformal: dict[str, dict[str, float]]) -> None:
    add_heading(doc, "1 课题主要研究内容及进度情况", 1)
    add_heading(doc, "1.1 主要研究内容", 2)
    add_body_paragraph(
        doc,
        "事件知识图谱是一类以事件为中心的结构化知识表示，它把文本中描述的事件抽取为节点，把事件之间的共指、时序、因果与子事件等关系组织为边，并在带时间戳的图结构上刻画事件的演化过程。与以实体为中心的传统知识图谱相比，事件知识图谱更强调“谁在何时对何对象做了何事”这一动态语义，因而适用于舆情分析、风险预警、政策追踪与未来态势研判等需要把握动态演化的场景。随着大规模语言模型在文本理解与生成上的发展，如何从中文长文本中自动构建事件知识图谱，并在其上进行可追溯、可控风险的时序推理，逐渐成为自然语言处理与知识工程交叉领域的重要问题。",
    )
    add_body_paragraph(
        doc,
        "从文本到可用图谱、再到可信预测，需要依次经过事件抽取、关系抽取与构图、时序推理三个环节。现有研究通常分别优化各环节的准确率指标，但抽取得到的论元是否有原文支撑、关系边是否接地且彼此一致、预测结果是否具备可追溯依据和错误率控制，仍难以在同一框架下检验。可检验性的不足会限制事件知识图谱在高风险决策场景中的使用。本课题据此研究一套以约束引导贯穿三个环节的事件知识图谱构建与时序推理方法，使事件节点、关系与预测分别在证据可回溯、约束可检验和错误率可控制三个方面具备可验证性。",
    )
    add_body_paragraph(
        doc,
        "需要说明的是，本课题的方法面向中文事件知识图谱的一般构建与时序推理问题，并不局限于某一特定领域。论文标题中不以金融作为研究对象限定，正文中所使用的ChFinAnn、DuEE-Fin、自建中文公告图谱以及FinDKG等数据，主要承担验证语料和对外比较的作用。选择中文金融公告作为主要验证域，是因为该类语料事件类型相对规范、时间线索明确，且有公开标准数据集便于复现实验；但方法本身围绕“中文事件文本的抽取、构图与推理”展开，原则上可迁移到新闻、公告、百科等其他中文事件语料。",
    )
    add_body_paragraph(
        doc,
        "围绕上述目标，学位论文拟展开三个相互衔接的研究内容：第一部分从非结构化中文文本中抽取带证据的事件节点；第二部分在事件节点之间建立证据接地、全局协调的关系边，形成事件图谱；第三部分在事件图谱或公开时序知识图谱上进行未来事件预测，并用保形预测控制分布漂移下的错误率。三部分的关系是“节点生成—关系组织—时序推理”的连续链条，而不是彼此割裂的模型堆叠。具体研究内容如下。",
    )
    add_heading(doc, "1.1.1 基于 SARGE 的文档级事件抽取", 3)
    add_body_paragraph(
        doc,
        "文档级事件抽取旨在把一篇完整文本转化为结构化的事件表，系统需要识别事件类型，并按照预定义角色填充主体、时间、数值、对象等论元。与句子级抽取相比，文档级抽取需要整合跨句分散的证据，还要处理一篇文档中多个同类或异类事件相互交织的情况；与此同时，下游图谱构建要求每个论元都能够回溯到原文片段。现有方法多依赖实体候选、图结构或代理节点等专用模块，流程较复杂，且证据链不易在输出中显式保留。基于此，本研究提出schema约束、角色接地的文档级事件抽取方法SARGE，把抽取重述为受schema约束的事件表生成任务，使模型直接生成带固定事件类型与角色槽位的结构化输出，并在后处理阶段进行模式检查和证据回填。该部分以ChFinAnn、DuEE-Fin和自建中文公告语料作为验证语料，目标是产出可供后续构图使用的带证据事件节点。",
    )
    add_heading(doc, "1.1.2 基于 GRPO-RLVR 的事件关系抽取与图谱构建", 3)
    add_body_paragraph(
        doc,
        "仅有孤立事件节点尚不足以支撑推理，还需要在节点之间建立可信的关系边。事件关系抽取的难点在于，单条边不仅要判断关系是否成立，还要尽量接地到文本证据，并满足共指等价、时序传递等全局一致性约束；若只进行逐对分类，容易得到局部正确但整体冲突的关系集合。基于此，本研究采用基于可验证奖励的生成式关系抽取思路，先用监督微调学习从事件窗口到关系集合的基本格式，再用GRPO-RLVR把格式合法、证据接地、全局一致与任务指标合成为统一奖励，引导模型抽取共指、时序、因果与子事件关系。抽取得到的局部关系再经一致性求解和闭包处理，形成证据接地、约束协调的事件图谱。",
    )
    add_heading(doc, "1.1.3 基于路径推理与保形预测的时序知识图谱推理", 3)
    add_body_paragraph(
        doc,
        "在构建好的事件图谱或公开时序知识图谱之上，时序推理负责根据历史事件预测未来可能发生的对象或关系。已有时序知识图谱推理方法通常以排序准确率为主要目标，较少同时考虑预测依据是否可追溯，以及在数据分布发生漂移时错误率是否可控。基于此，本研究把路径策略强化学习、强基线融合与保形预测结合起来：一方面沿历史图路径搜索候选答案，使预测能够回溯到历史证据链；另一方面用保形预测把单点排序转化为带覆盖保证的预测集合，从而为分布漂移下的预测结果提供可量化的风险控制。",
    )
    add_body_paragraph(
        doc,
        "上述三部分由两条主线贯穿为整体。其一是统一的事件表示：第一部分抽取得到的事件节点既是第二部分关系构建的输入，也是第三部分时序推理的状态基础，同一批数据能够经抽取、构图、推理逐级贯通。其二是统一的约束引导机制：同一可验证性原则在抽取阶段表现为证据门控，在关系阶段表现为训练奖励，在推理阶段表现为风险控制器。所谓门控、奖励与风险控制并非三个割裂模块，而是同一约束思想在不同研究阶段的具体化。方法验证采用公开基准与自建中文事件图谱两层数据：公开基准承担对外可比，自建图谱承担三阶段接口贯通的闭环演示。",
    )
    add_heading(doc, "1.2 进度情况", 2)
    add_body_paragraph(
        doc,
        "截至中期，三部分研究内容均已跑通主干并取得可复现的阶段性结果，三部分数据接口已贯通为闭环，总体进度约过半。事件抽取部分已完成文档级事件抽取的公开基线对比与事件节点转换，触发词跨度与时间证据的稳定导出尚在补齐；关系抽取与图谱构建部分已完成监督微调与 GRPO-RLVR 的对照实验与一致性事件图谱原型，学习式中文关系迁移与边质量评估有待开展；时序推理部分已完成路径推理、强基线融合与保形预测风险控制，多种子、多数据集消融与下游应用尚未完成。当前阶段的定位是在可比口径下接近经典强基线，并稳定提供纯准确率方法所缺少的证据链与覆盖保证，尚未声称在全部公开基准上超过最新方法。三部分内容的完成情况、关键结果与待办事项概括于表1。",
    )
    add_table(
        doc,
        ["研究内容", "已完成进展", "当前最关键证据", "后续升级方向"],
        [
            ["一 SARGE 事件抽取", "文档级事件表生成、公开基线对比、节点转换", "ChFinAnn F1 86.0；DuEE-Fin F1 78.0", "触发词/时间证据、记录绑定"],
            ["二 关系抽取与构图", "监督微调→GRPO-RLVR、证据接地、一致性图原型", "CoNLL 共指 0.265→0.771；时序 P 0.445→0.565", "学习式中文关系迁移、边质量抽样"],
            ["三 时序推理与风控", "Path-RL、RE-GCN、融合搜索、保形预测", f"融合搜索 MRR {ch3['fusion_sweep']['mrr_tfilt']:.3f}；ACI 覆盖率 {conformal['aci']['coverage']:.3f}", "多数据集、多种子、消融与下游选择性预测"],
        ],
        "表1 三部分研究内容中期进展总表",
        widths_cm=[2.7, 4.4, 4.1, 4.0],
    )


def add_results_section(doc: Document, ch3: dict[str, dict[str, float | str]], conformal: dict[str, dict[str, float]]) -> None:
    add_heading(doc, "2 目前已完成的研究工作及结果", 1)
    add_heading(doc, "2.1 总体技术路线与统一数据表示", 2)
    add_body_paragraph(
        doc,
        "本阶段已完成的工作按照“数据处理与统一表示—模型设计—实验设置—结果对比—误差诊断”的链条展开。首先，事件抽取从中文长文本中生成受schema约束的事件表，并将其规范化为带证据的事件节点；其次，关系抽取与图谱构建在事件节点之间建立共指、时序、因果与子事件关系，并通过证据接地与一致性求解去除不可验证的边；最后，时序推理把事件图谱或公开时序知识图谱表示为带时间戳的四元组，在其上完成未来查询的实体排序与风险控制。这样，事件节点同时充当抽取产物、构图输入和推理状态，三个部分在同一数据对象上接续推进。"
    )
    add_body_paragraph(
        doc,
        "在这一统一表示之上，本研究把可验证性落实为贯穿三阶段的验证器机制。抽取阶段要求论元能够回溯到原文证据，构图阶段要求关系边接地并与全局约束相容，推理阶段要求预测路径有据可循且预测集合在分布漂移下满足覆盖要求。三者本质上都可归结为对候选输出施加可检验约束并据此取舍：当约束用于候选论元过滤时，验证器表现为门控；当约束转化为训练反馈时，验证器表现为奖励；当约束用于部署阶段的错误率控制时，验证器表现为风险控制器。上述数据流与验证器主线的整体关系如图1所示。",
    )
    add_figure(doc, FIGURES / "fig_framework.png", "图1 事件知识图谱构建与时序推理总体框架", 16.5)
    add_body_paragraph(
        doc,
        "为统一说明三个部分的实验基础，表2汇总了本课题已用与计划使用的数据集。ChFinAnn、DuEE-Fin、MAVEN-ERE和ICEWS14等公开基准用于提供对外可比的指标；自建中文事件图谱event_graph_zh用于衔接三部分数据接口并演示从文本到图谱再到推理的闭环；标注[计划]的数据集用于后续验证域扩展、漂移鲁棒性评测和下游应用。除自建图谱外，公开数据集均沿用原论文划分，不重新切分，以保证评测口径可比。需要强调的是，表中金融相关数据集主要承担验证作用，并不意味着方法仅面向金融事件。",
    )
    add_table(
        doc,
        ["数据集", "用途", "规模（训练/验证/测试）", "官方划分", "备注"],
        [
            ["ChFinAnn", "文档级事件抽取（一）", "公开标准划分", "是", "SARGE 主评测之一"],
            ["DuEE-Fin", "文档级事件抽取（一）", "公开标准划分，dev500 评测", "是", "SARGE 主评测之一"],
            ["MAVEN-ERE", "事件关系抽取（二）", "2913 / 710 / 857", "是", "test 标签隐藏，按 valid 评测"],
            ["CCKS-2021 因果", "金融因果抽取（二）", "5600 / 700 / 700", "本地按 seed 划分", "官方无公开 test；计划补充"],
            ["ICEWS14", "时序图谱预测（三）", "63685 / 13823 / 13222", "是（TiRGN 划分）", "时间外推、三段零重叠"],
            ["ICEWS18 / 05-15", "漂移扩展（三）", "37 万级，两个漂移集", "是", "风控漂移强化；计划"],
            ["FinDKG", "金融时序预测（三）", "119549 / 11444 / 13069", "是", "金融域对标；计划"],
            ["Astock / CMIN-CN", "下游涨跌预测（三）", "Astock 11815/1477/1477；CMIN 300 股", "是", "下游应用；计划"],
            ["event_graph_zh（自建）", "闭环演示", "677 节点 / 497 证据边 / 20683 总边", "自建", "无外部基线，闭环载体"],
        ],
        "表2 数据集与评测口径清单",
        widths_cm=[2.8, 3.5, 4.3, 1.9, 3.0],
    )

    add_heading(doc, "2.2 基于 SARGE 的文档级事件抽取", 2)
    add_body_paragraph(
        doc,
        "文档级事件抽取是整条链路的入口，其目标是把长文本中的事件类型、角色和值组织为结构化事件表。该任务的主要困难在于，论元常跨句分散，一篇文档中又可能包含多条同类型记录，角色取值恢复和事件行绑定都容易出错。为解决这一问题，本节提出的SARGE把抽取任务改写为schema约束下的事件表生成：模型不生成自由文本，而是在给定事件类型和角色槽位约束下直接输出固定槽事件行，再通过模式检查、固定槽规范化和证据回填形成可用于构图的事件节点。方法以Qwen3-4B-Instruct-2507为基座，采用4-bit LoRA/QLoRA进行监督微调；评测口径为字段级micro平均F1，即在固定schema角色槽位上逐字段计算精确率、召回率和F1。"
    )
    add_body_paragraph(
        doc,
        "为验证方法有效性，本节在ChFinAnn和DuEE-Fin两个公开中文文档级事件抽取数据集上与Doc2EDAG、GIT、EPAL和SEELE等公开基线比较。表3给出字段级F1结果，图2将其可视化。由表3可知，在ChFinAnn上，SARGE取得86.0的F1，高于Doc2EDAG、GIT、EPAL并略优于SEELE；在DuEE-Fin上，SARGE取得78.0的F1，略低于SEELE的80.8，但高于其余基线。综合来看，生成式事件表方法在传统专用监督模型长期占优的文档级抽取任务上具有竞争力，并且在两个公开基准上均处于第一梯队。更值得关注的是多事件场景：SARGE在该子集上的F1为77.5，相对单事件场景的回落幅度较小，说明结构化生成对跨句论元和事件交织具有一定稳健性。三次随机种子重复实验的标准差分别为±0.39和±0.38，结果较稳定。",
    )
    add_table(
        doc,
        ["数据集", "Doc2EDAG", "GIT", "EPAL", "SEELE", "SARGE"],
        [
            ["ChFinAnn", "78.8", "80.3", "83.4", "85.1", "86.0"],
            ["DuEE-Fin", "63.4", "67.8", "76.4", "80.8", "78.0"],
        ],
        "表3 SARGE与公开基线对比",
        widths_cm=[3.0, 2.4, 2.2, 2.2, 2.2, 2.2],
    )
    add_figure(doc, FIGURES / "fig_ch1_sarge_baselines.png", "图2 本文方法与公开基线在中文文档级事件抽取上的性能对比", 14.5)
    add_body_paragraph(
        doc,
        "SARGE的输出不是最终目标，而是图谱构建的节点入口，因此还需要考察其规范化后的结构特征。表4汇总了自建中文事件节点统计：同一批中文公告经SARGE抽取并规范化后，共得到677个事件节点，覆盖429篇文档和13类事件类型，平均每个事件携带5.02个结构化论元，主体覆盖率达到100%（478个不同主体）。这些结果说明当前节点侧的事件类型与论元信息已能支撑事件中心图结构的构建。与此同时，时间锚点覆盖率仅为30%，且规范化预测尚未稳定导出触发词跨度，目前只能从原文回填论元级证据。也就是说，第一部分已经完成了高质量事件节点的主体工作，但触发词、时间证据和记录级绑定仍是后续需要补齐的节点接口问题。",
    )
    add_table(
        doc,
        ["指标", "数值"],
        [
            ["事件节点数", "677"],
            ["来源文档数", "429"],
            ["事件类型数", "13"],
            ["平均论元数", "5.02"],
            ["时间锚点覆盖率", "30%"],
            ["主体覆盖率", "100%（478 个主体）"],
            ["论元证据覆盖", "673/677 节点有回填证据"],
        ],
        "表4 SARGE自建中文事件节点统计",
        widths_cm=[5.0, 8.0],
    )
    add_body_paragraph(
        doc,
        "图3进一步给出这677个事件节点的类型分布。由图可知，股份回购、质押、亏损、中标等高频类型构成主体，类型分布虽不均衡，但已覆盖股权变动、经营业绩与监管处罚等多类事件，为后续事件图谱提供了类型多样的节点来源。少数低频类型样本偏少，说明后续扩充验证语料时需要继续补强长尾事件。",
    )
    add_figure(doc, FIGURES / "fig_ch1_event_types.png", "图3 自建中文事件图谱的事件类型分布", 12.8)

    add_heading(doc, "2.3 基于 GRPO-RLVR 的事件关系抽取与图谱构建", 2)
    add_body_paragraph(
        doc,
        "在获得事件节点之后，第二部分需要把离散节点组织为可推理的事件图谱。该环节的难点不只在于判断某一对事件之间是否存在关系，还在于关系边需要尽量接地到文本证据，并在全局上满足共指等价、时序传递等一致性约束。若只进行逐对分类，模型可能得到局部合理但整体互相冲突的边集合。基于此，本节采用生成式关系抽取与可验证奖励相结合的思路：先用监督微调使模型掌握从事件窗口到关系集合的基本输出格式，再用GRPO-RLVR把格式合法、证据接地、全局一致与任务F1合成为训练奖励，使可验证约束在训练阶段就参与模型优化。关系类型涵盖共指、时序、因果与子事件四类。",
    )
    add_body_paragraph(
        doc,
        "实验首先在MAVEN-ERE验证集上进行。该数据集统一标注事件共指、时序、因果与子事件关系，适合考察多关系类型下的抽取与一致性能力。表5列出监督微调与GRPO-RLVR的对比结果，图4进一步展示事件共指质量。由表5可知，加入可验证奖励后，微平均精度由0.332提升到0.647，说明模型输出的关系更少臆造；CoNLL共指F1由0.265提升到0.771，表明一致性奖励有助于把分散的指代判断收敛为协调的等价类。时序关系仅报告精度，是因为MAVEN-ERE的时序标注包含稠密传递闭包，而生成式抽取输出天然稀疏，直接比较召回率或F1会混入标注密度差异。需要说明的是，微平均F1的绝对数值仍然偏低，这与生成式稀疏抽取和监督式事件对分类的评测口径尚未完全对齐有关，因此当前结果主要证明可验证奖励相对监督微调的改进，而不声称已超过监督式强基线。",
    )
    add_table(
        doc,
        ["指标", "监督微调", "GRPO-RLVR"],
        [
            ["微平均 F1", "0.012", "0.070"],
            ["微平均精度", "0.332", "0.647"],
            ["共指 F1", "0.060", "0.611"],
            ["CoNLL 共指 F1", "0.265", "0.771"],
            ["时序关系精度", "0.445", "0.565"],
        ],
        "表5 监督微调与 GRPO-RLVR 关系抽取结果",
        widths_cm=[5.0, 4.0, 4.0],
    )
    add_figure(doc, FIGURES / "fig_coref_conll.png", "图4 监督微调与 GRPO-RLVR 的事件共指质量对比", 12.2)
    add_body_paragraph(
        doc,
        "在公开基准之外，本节还把事件节点和关系约束应用到自建中文验证语料，形成event_graph_zh图谱原型，其整体概览见图5。该图把677个事件节点按13类事件着色，并画出经证据接地保留的关系骨架，其中实线为时序边、虚线为共指边。由图可见，节点大体按公告和主体形成局部簇，同一篇文档内的事件经共指与时序关系连成小图，而簇间相对稀疏；图中约251个节点参与了证据关系，其余节点暂为孤立事件，说明证据边覆盖仍有提升空间。这里的自建图谱用于演示三部分接口贯通，尚不能等同于已经完成的学习式中文关系抽取结果。",
    )
    add_figure(doc, FIGURES / "fig_event_graph_overview.png", "图5 自建中文事件图谱整体概览", 16.5)
    add_body_paragraph(
        doc,
        "图6从数量角度展示图谱构建与一致性闭包过程。在677个事件节点上，启发式抽取产生498条候选关系边，其中1条因缺乏证据被过滤，最终保留497条证据候选边；随后经一致性求解与时序、共指闭包，扩展为20683条图边。由图6可知，闭包使边数扩张约41倍，说明两万余条图边中大多数来自共指等价类展开和时序传递闭包，而非独立抽取的原始证据边。因此，当前中文图谱更准确的定位是“证据接地的图谱原型和闭环载体”，后续还需用学习式关系抽取替换启发式关系来源，并系统评估闭包误差传播。",
    )
    add_figure(doc, FIGURES / "fig_event_graph_funnel.png", "图6 中文事件图谱的构建与一致性闭包过程", 14.5)
    add_body_paragraph(
        doc,
        "图7进一步放大一个真实局部子图，展示事件中心知识图谱的微观形态。圆形节点表示事件，角色边连接标的公司、主体、时间和数值等取值，事件之间则由先于关系和共指关系相连。该例说明，同一主体的多起同类事件可以被组织为可追溯的时间链，从而为后续路径推理提供结构基础。综合表5和图5至图7可以看出，第二部分已经完成了公开基准上的GRPO-RLVR对照实验，并构建了自建中文事件图谱原型；但中文图谱的关系质量评估、学习式迁移和闭包误差控制仍需在后续阶段补齐。",
    )
    add_figure(doc, FIGURES / "fig_event_graph_example.png", "图7 新奥股份局部事件中心知识图谱实例", 17.0)

    add_heading(doc, "2.4 基于路径推理与保形预测的时序知识图谱推理", 2)
    add_body_paragraph(
        doc,
        "第三部分面向事件图谱上的未来预测问题。已有时序知识图谱推理方法通常直接优化候选实体排序准确率，但在中期报告所关注的完整研究链条中，预测结果还需要满足两项额外要求：一是能够回溯到历史证据路径，二是在数据分布漂移时具有可量化的错误率控制。为此，本节将路径策略强化学习、强基线融合和保形预测结合起来。给定历史四元组，模型输出未来查询下的候选实体排序，并通过路径搜索提供可解释证据；在排序分数之上，再用保形预测构造带覆盖保证的候选集合。中期实验以ICEWS14为公开基准，采用时间感知过滤的MRR与Hits作为主口径，以严格区分面向未来的外推预测和同时刻补全任务。",
    )
    add_body_paragraph(
        doc,
        f"表6与图8汇总了ICEWS14上从基线到融合搜索的阶段性结果。早期Path-RL的MRR为{ch3['path_rl']['mrr']:.3f}，高于频率基线和简化时序GNN基线；随后引入复发基线、RE-GCN强基线和融合权重搜索。其中，RE-GCN的时间感知过滤MRR为{ch3['re_gcn']['mrr_tfilt']:.3f}，单次融合为{ch3['hybrid_single']['mrr_tfilt']:.3f}，在{ch3['fusion_sweep']['weights']}下经权重搜索达到{ch3['fusion_sweep']['mrr_tfilt']:.3f}。由表6可知，从频率下界到路径推理、再到强基线和融合搜索，指标逐级提升，说明局部路径信息与全局历史模式具有互补性。需要说明的是，{ch3['fusion_sweep']['mrr_tfilt']:.3f}是公开基准上经融合权重搜索得到的阶段性最优结果，并非单一模型的最终指标；完整的RE-GCN扩展、多种子实验和机制消融仍在推进。",
    )
    add_table(
        doc,
        ["方法", "MRR", "Hits@1", "Hits@10", "说明"],
        [
            ["Frequency", f"{ch3['frequency']['mrr']:.3f}", f"{ch3['frequency']['hits@1']:.3f}", f"{ch3['frequency']['hits@10']:.3f}", "朴素频率下界"],
            ["Temporal-GNN", f"{ch3['temporal_gnn']['mrr']:.3f}", f"{ch3['temporal_gnn']['hits@1']:.3f}", f"{ch3['temporal_gnn']['hits@10']:.3f}", "简化神经基线"],
            ["Recurrency", f"{ch3['recurrency']['mrr']:.3f}", f"{ch3['recurrency']['hits@1']:.3f}", f"{ch3['recurrency']['hits@10']:.3f}", "多信号 copy"],
            ["Path-RL", f"{ch3['path_rl']['mrr']:.3f}", f"{ch3['path_rl']['hits@1']:.3f}", f"{ch3['path_rl']['hits@10']:.3f}", "路径策略"],
            ["RE-GCN", f"{ch3['re_gcn']['mrr_tfilt']:.3f}", f"{ch3['re_gcn']['hits@1_tfilt']:.3f}", f"{ch3['re_gcn']['hits@10_tfilt']:.3f}", "强基线"],
            ["融合搜索", f"{ch3['fusion_sweep']['mrr_tfilt']:.3f}", f"{ch3['fusion_sweep']['hits@1_tfilt']:.3f}", f"{ch3['fusion_sweep']['hits@10_tfilt']:.3f}", str(ch3["fusion_sweep"]["weights"])],
        ],
        "表6 ICEWS14 时序预测结果",
        widths_cm=[3.1, 2.1, 2.1, 2.1, 5.2],
    )
    add_figure(doc, FIGURES / "fig_forecasting_compare.png", "图8 ICEWS14 上各方法的时序预测性能", 14.2)
    add_body_paragraph(
        doc,
        "图9展示Path-RL的训练过程。平均奖励与命中率随训练轮次上升并逐渐趋于平稳，表明路径策略能够在当前设定下稳定学习。结合表6可以看出，路径推理不仅提供了可解释的历史证据链，也为后续与强基线融合提供了有效的局部推理信号。",
    )
    add_figure(doc, FIGURES / "fig_pathrl_training.png", "图9 Path-RL 的训练过程", 12.2)
    add_body_paragraph(
        doc,
        "为说明当前结果在公开基准上的位置，表7将本文与近年代表性方法在同一口径下进行对照。本文当前最优MRR为0.411，与RE-GCN报告的约0.42基本接近，较TiRGN的0.440和L2TKG的0.459仍有差距。该差距需要在后续通过更完整的强基线训练、多数据集实验和消融分析继续缩小；但与此同时，这些以准确率为主的代表方法通常不提供下文所讨论的漂移鲁棒覆盖能力。",
    )
    add_table(
        doc,
        ["方法", "年份", "MRR（过滤）", "口径与说明"],
        [
            ["本文 融合搜索", "2026", "0.411", "本文当前最优"],
            ["TITer", "2021", "≈0.41", "路径强化学习，与本文同类"],
            ["RE-GCN", "2021", "≈0.42", "演化式 GCN 强基线"],
            ["TiRGN", "2022", "0.440", "时间引导的局部-全局历史"],
            ["L2TKG", "2023", "0.459", "学习潜在关系，近年代表"],
        ],
        "表7 ICEWS14 与近年方法口径对照",
        widths_cm=[3.0, 1.8, 2.6, 6.4],
    )
    add_body_paragraph(
        doc,
        "在排序准确率之外，本节进一步验证风险控制能力，这对应2.1节所述验证器的风险控制身份。保形预测的作用不是改变排序模型本身，而是在排序分数之上构造带覆盖保证的预测集合。表8与图10比较了静态划分、自适应、近因加权和CRC四类校准策略。在目标覆盖率0.9下，静态划分校准在分布漂移时的实际覆盖率为0.858，漂移覆盖失真为0.29；自适应校准的实际覆盖率达到0.899，漂移覆盖失真降至0.20。由此可见，当数据分布发生变化时，自适应校准能够更稳定地把覆盖率维持在目标附近，从而为下游决策提供可量化的错误率保证。需要同时看到，自适应校准的平均集合大小更大，说明覆盖率提升需要以候选集合扩张为代价。",
    )
    add_table(
        doc,
        ["校准器", "覆盖率", "漂移覆盖失真↓", "平均集合大小"],
        [
            ["split", f"{conformal['split']['coverage']:.3f}", f"{conformal['split']['drift_gap']:.2f}", f"{conformal['split']['set_size']:.1f}"],
            ["ACI", f"{conformal['aci']['coverage']:.3f}", f"{conformal['aci']['drift_gap']:.2f}", f"{conformal['aci']['set_size']:.1f}"],
            ["weighted", f"{conformal['weighted']['coverage']:.3f}", f"{conformal['weighted']['drift_gap']:.2f}", f"{conformal['weighted']['set_size']:.1f}"],
            ["CRC", f"{conformal['crc']['coverage']:.3f}", f"{conformal['crc']['drift_gap']:.2f}", f"{conformal['crc']['set_size']:.1f}"],
        ],
        "表8 保形预测风险控制结果",
        widths_cm=[3.3, 3.0, 3.0, 4.0],
    )
    add_figure(doc, FIGURES / "fig_conformal_risk_control.png", "图10 不同校准策略下的覆盖率与漂移失真", 14.0)

    add_body_paragraph(
        doc,
        "综观三个部分的中期结果，第一部分已经完成schema约束、角色接地的文档级事件表生成，并在两个公开中文数据集上取得竞争性结果；第二部分已经完成MAVEN-ERE上的SFT与GRPO-RLVR对照，并构建出自建中文事件图谱原型；第三部分已经完成ICEWS14上的路径推理、强基线融合和保形预测风险控制。三部分之所以能够连成整体，是因为抽取得到的事件节点向后进入关系构建，关系图谱再向后支撑时序推理；验证器也随之从门控、奖励到风险控制依次承担不同角色。中期阶段不宣称在全部公开基准上超过最新方法，而是在可比口径下接近经典强基线，并补充准确率导向方法较少提供的证据链与覆盖保证。"
    )


def add_problem_section(doc: Document) -> None:
    add_heading(doc, "3 目前存在的或预期可能出现的问题", 1)
    paragraphs = [
        "在本阶段研究工作中，尽管三部分研究内容均已跑通主干并取得可复现的阶段性结果，三部分数据接口也已贯通为闭环，但从完整学位论文的要求来看，仍存在或预期可能出现若干需要进一步解决的问题。",
        "（1）从文本到图谱的字段契约仍不完整。当前SARGE已经能够稳定输出事件类型和主要论元，但触发词跨度尚未稳定导出，时间锚点覆盖率仅为30%，同质记录绑定问题也尚未解决。这意味着事件节点虽然已经可以用于构图原型，但在时间证据、触发词证据和整行记录组织方面仍不够完整，后续需要从输出格式、时间表达式识别、记录级绑定约束和证据字段校验等方面补齐。",
        "（2）自建中文事件图谱仍以启发式关系为主。event_graph_zh已经证明了从事件节点到事件中心图结构的可构建性，但其两万余条边中大量来自共指等价类展开与时序传递闭包，原始证据边数量相对有限，且中文关系尚未全面迁移为学习式抽取结果。因此，当前图谱更适合作为闭环演示和接口验证载体，尚不能作为最终的学习式中文事件关系图谱。后续需要开展边抽样准确率评估、闭包误差传播分析和中文关系迁移实验。",
        "（3）评测与消融的完备性仍需加强。关系抽取部分与监督式强基线的任务口径尚未完全对齐，需要补充同设定事件对分类底座，或更清晰地区分生成式稀疏抽取与监督式稠密分类的差异；同时还需补全证据接地、一致性约束和任务奖励三者的消融。时序推理部分虽已完成ICEWS14上的路径推理、RE-GCN和融合搜索，但多种子、多数据集以及路径可信度奖励、势函数塑形、热启动和组内相对基线等机制消融仍需进一步开展。",
        "（4）风险控制与下游应用的连接仍然偏弱。自适应保形校准在ICEWS14上已经表现出更稳定的覆盖率，但选择性预测、覆盖率约束下的准确率以及事件驱动下游任务仍未形成完整实验链条。后续需要在更密的中文事件新闻图谱或下游预测数据集上进一步验证，使风险控制不只停留在公开时序基准上，而能与实际应用场景相连接。",
        "总体来看，上述问题主要属于把已跑通主干继续做扎实、做完整的工作，而不是研究路线层面的根本障碍。下一阶段的重点是补全证据字段、升级关系来源、完善评测消融，并把风险控制连接到下游任务，使当前阶段性结论进一步发展为论文级证据。",
    ]
    for paragraph in paragraphs:
        add_body_paragraph(doc, paragraph)


def add_schedule_section(doc: Document) -> None:
    add_heading(doc, "4 后期学术研究的进度安排", 1)
    add_heading(doc, "4.1 拟完成的研究工作", 2)
    add_body_paragraph(
        doc,
        "为克服当前存在或预期可能出现的问题，并按时完成整个课题目标，后续研究将围绕“补证据、强图谱、全评测、连下游”四条线索展开。相关工作与第3部分指出的问题一一对应，整体遵循先补齐节点证据、再增强关系图谱、进而完善推理评测和下游验证的顺序推进。",
    )
    add_body_paragraph(
        doc,
        "（1）补齐SARGE的证据字段并改进记录绑定。后续将在事件表输出契约中显式加入触发词槽位和规范化时间字段，引入面向同质记录的去重、对齐和记录级约束，并在ChFinAnn、DuEE-Fin及自建中文语料上以字段级F1和精确记录匹配率为主要指标进行系统评估。该项工作的目标是把当前较高的角色级抽取能力进一步转化为更可靠的事件节点质量。",
    )
    add_body_paragraph(
        doc,
        "（2）开展学习式中文事件关系抽取与图谱质量评估。后续将把当前证据接地的启发式关系图升级为学习式关系抽取结果，补充边抽样准确率评估、闭包误差传播分析和中文关系迁移实验，并加入同设定事件对分类底座与GRPO奖励消融。该项工作的目标是让自建中文图谱由“可构图原型”进一步转化为质量可度量、来源可复现的关系图谱。",
    )
    add_body_paragraph(
        doc,
        "（3）完善时序推理评测并连接下游应用。后续将补全Path-RL与RE-GCN的多种子实验和机制消融，把公开基准从ICEWS14扩展到ICEWS18、ICEWS05-15和FinDKG，并进一步验证自适应、近因加权和CRC等风险控制策略。在条件允许的情况下，还将结合选择性预测和下游事件驱动任务，考察覆盖率约束下的实际决策效果。",
    )
    add_body_paragraph(
        doc,
        "上述三项工作之间具有明确的承接关系：第一项为第二项提供更完整的事件节点输入，第二项为第三项提供质量可度量的图谱底座，第三项则把整条链路的价值落实到可追溯、可控风险的预测与下游应用上。预期通过这些工作，可以在保证三部分闭环与可复现的基础上，把当前“接近强基线、具备可验证性”的阶段性结果进一步充实为完整、系统的学位论文成果。",
    )
    add_heading(doc, "4.2 进度安排", 2)
    add_body_paragraph(
        doc,
        "为保证上述工作按时完成，本研究对后续阶段作出如下时间规划。总体顺序是先补齐证据字段，再增强中文图谱，进而完善公开基准评测和风险控制实验，最后完成论文整合与预答辩准备。各阶段时间、主要任务与预期产出如表9所示。",
    )
    add_table(
        doc,
        ["时间", "主要任务", "预期产出"],
        [
            ["2026年7月—9月", "补齐 SARGE 触发词与时间证据；整理事件抽取错误分析与记录绑定诊断", "事件抽取定稿表、图与错误分析"],
            ["2026年9月—11月", "中文学习式关系图迁移；边抽样准确率评估；GRPO 奖励消融", "关系抽取与图谱构建完整实验矩阵与中文图质量报告"],
            ["2026年11月—2027年1月", "ICEWS18、ICEWS05-15、FinDKG 扩展；Path-RL 与 RE-GCN 多种子消融", "时序推理公开基准表、训练曲线与路径案例"],
            ["2027年1月—3月", "自适应、近因加权与风险控制（ACI、weighted、CRC）扩展；选择性预测与下游交易评估", "风险控制补充实验与应用讨论"],
            ["2027年3月—4月", "整合全文、统一术语与引用，完善复现说明，准备预答辩", "学位论文初稿、答辩幻灯片、代码与数据说明"],
        ],
        "表9 后续研究进度安排",
        widths_cm=[3.6, 7.0, 4.8],
    )





def normalize_sections(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def main() -> None:
    ch3 = load_ch3_results()
    conformal = load_conformal_results()

    shutil.copyfile(TEMPLATE, OUT_DOCX)
    doc = Document(OUT_DOCX)
    normalize_sections(doc)
    replace_title_and_clear_body(doc)  # 改标题 + 删旧正文，保留封面
    add_progress_section(doc, ch3, conformal)
    add_results_section(doc, ch3, conformal)
    add_problem_section(doc)
    add_schedule_section(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
