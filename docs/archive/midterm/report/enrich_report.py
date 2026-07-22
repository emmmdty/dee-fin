#!/usr/bin/env python
"""在 `童佳锴中期报告初稿_答辩版.docx` 上原地增补：
  1) 强化"实际场景 / 技术价值"论述（外行懂 + 内行认可）；
  2) 把 22 条核心公式以 Word 原生公式（pandoc→OMML，编号 (2-x)）插入 2.1~2.4；
  3) 4 段算法伪代码框（SARGE / GRPO-RLVR / Path-RL / 保形风控）；
  4) 用 DNN 风格分层架构图替换图2-1。

依赖：pandoc、python-docx（项目 venv 已含）。公式来源 = docs/midterm/report/中期报告公式补充.md，
均与 src/finekg 下代码一致。运行：
    .venv/bin/python docs/midterm/report/enrich_report.py
之后用 soffice 转 PDF。
"""
from __future__ import annotations

import copy
import shutil
import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

HERE = Path(__file__).parent
REPORT = HERE / "童佳锴中期报告初稿_答辩版.docx"
BASE_REPORT = HERE / "童佳锴中期报告初稿_答辩版.bak.docx"
ARCH_PNG = HERE.parent / "figures" / "fig_architecture.png"
ARCH_DETAIL_PNG = HERE.parent / "figures" / "fig_architecture_detail.png"
# ----------------------------- 低层工具 -----------------------------
def insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def copy_ppr(src: Paragraph, dst: Paragraph) -> None:
    s = src._p.find(qn("w:pPr"))
    if s is None:
        return
    d = dst._p.find(qn("w:pPr"))
    if d is not None:
        dst._p.remove(d)
    dst._p.insert(0, copy.deepcopy(s))


def copy_rpr(src_run, dst_run) -> None:
    s = src_run._r.find(qn("w:rPr"))
    if s is None:
        return
    d = dst_run._r.find(qn("w:rPr"))
    if d is not None:
        dst_run._r.remove(d)
    new = copy.deepcopy(s)
    # 显式设定 eastAsia=宋体、cs=宋体（中文标点：双引号、破折号、顿号等）
    rfonts = new.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        new.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "宋体")
    dst_run._r.insert(0, new)


def _ref_run(ref: Paragraph):
    for r in ref.runs:
        if r.text.strip():
            return r
    return ref.runs[0] if ref.runs else None


def set_mono(run, size=9.5) -> None:
    run.font.size = Pt(size)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Consolas")
    rfonts.set(qn("w:hAnsi"), "Consolas")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "宋体")


FORMULA_LINES = {
    "2-1": [
        "max_{M1,M2,M3} lambda1*F1_role + lambda2*F1_rel + lambda3*MRR - lambda4*Risk",
        "s.t. g(v,D)=1, g(e,D)=1, Consistent(E)=1, Pr[y in C(x)] >= 1-alpha",
    ],
    "2-2": ["Verifier = { Gate: extraction; Reward: relation; RiskController: inference }"],
    "2-3": ["e = (t, {(r, V_r, A_r) | r in R_t})"],
    "2-4": ["Y_hat_D = argmax_{Y in Y(D,S)} p_theta(Y | D, S, C_K(D), P_D)"],
    "2-5": ["L_SFT = - sum_{(D,Y*)} sum_j log p_theta(y*_{D,j} | x_D, y*_{D,<j})"],
    "2-6": [
        "P_role = TP/(TP+FP),   R_role = TP/(TP+FN)",
        "F1_role = 2*P_role*R_role/(P_role+R_role)",
    ],
    "2-7": [
        "ExactRec = 2*sum_{D,t} sum_{(e,e_hat) in M_{D,t}} 1[e=e_hat]",
        "           / sum_{D,t} (|Y_{D,t}| + |Y_hat_{D,t}|)",
        "Delta_bind = F1_role - ExactRec",
    ],
    "2-8": ["R(x,y)=w_f*R_format + w_g*R_ground + w_c*R_consist + w_t*R_task"],
    "2-9": [
        "R_format = 1[Parse(y)!=empty] * |ValidEdges(y)| / max(|Edges(y)|,1)",
        "R_ground = sum_{e in E_hat} 1[eta_e subset D] / max(|E_hat|,1)",
        "R_consist = 1 - |E_hat - Pi(E_hat)| / max(|E_hat|,1);  R_task = F1_rel(E_hat,E*)",
    ],
    "2-10": ["A_hat_i = (R_i - mean(R_1,...,R_G)) / max(std(R_1,...,R_G), epsilon)"],
    "2-11": [
        "J_GRPO(theta) = (1/G) * sum_i min(rho_i*A_hat_i, clip(rho_i,1-eps,1+eps)*A_hat_i)",
        "                - beta * KL(pi_theta || pi_ref)",
    ],
    "2-12": [
        "v_i ~ v_j => v_j ~ v_i;  (v_i ~ v_j and v_j ~ v_k) => v_i ~ v_k",
        "(v_i < v_j and v_j < v_k) => v_i < v_k;  E_temp+ = TransitiveClosure(E_temp)",
    ],
    "2-13": [
        "FNR_hat(tau) = (1/n) * sum_i 1[s_{e*_i} < tau]",
        "U_CRC(tau) = n/(n+1) * FNR_hat(tau) + 1/(n+1)",
        "tau* = max{tau : U_CRC(tau) <= alpha}",
    ],
    "2-14": ["f_theta(o | q, H_{<tau_q}),   H_{<tau_q}={(s,r,o,tau) in H | tau < tau_q}"],
    "2-15": [
        "s_l = (q, e_l, path_{0:l}),   e_0 = s_q",
        "A(s_l) = {(r,e',tau') : (e_l,r,e',tau') in H_{<tau_q}} union {STAY}",
    ],
    "2-16": ["R(xi)=c_h*1[e_L=o*] + c_f*F(xi)*1[e_L=o*] + c_s*sum_l(gamma*Phi(s_{l+1})-Phi(s_l))"],
    "2-17": ["F(xi)=max(0, f_theta(o*|q,H) - f_theta(o*|q,H without xi))"],
    "2-18": ["MRR=(1/N)*sum_i 1/rank_i,   Hits@k=(1/N)*sum_i 1[rank_i <= k]"],
    "2-19": ["rank_tfilt(o*) = 1 + sum_{o in V-F(q)} 1[f_theta(o|q) > f_theta(o*|q)]"],
    "2-20": ["C_qalpha(x)={o in V : rank(o|x) <= q_alpha},   q_alpha=a_(ceil((n+1)(1-alpha)))"],
    "2-21": ["alpha_{t+1}=alpha_t + gamma*(alpha-err_t),   err_t=1[y_t notin C_t(x_t)]"],
    "2-22": [
        "Cov_hat=(1/N)*sum_t 1[y_t in C_t(x_t)]",
        "Size_hat=(1/N)*sum_t |C_t(x_t)|",
        "DriftGap=max_{W in windows} | mean_{t in W} 1[y_t in C_t(x_t)] - (1-alpha) |",
    ],
}


def set_formula_run(run, size: float = 10.2) -> None:
    run.font.size = Pt(size)
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:cs"), "宋体")


def clear_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))





# ----------------------------- 块插入器 -----------------------------
def insert_body(cursor: Paragraph, text: str, ref: Paragraph) -> Paragraph:
    p = insert_paragraph_after(cursor)
    copy_ppr(ref, p)
    run = p.add_run(text)
    rr = _ref_run(ref)
    if rr is not None:
        copy_rpr(rr, run)
    return p


def insert_equation(cursor: Paragraph, latex: str, num: str, usable_emu: int) -> Paragraph:
    """Insert a LibreOffice-safe equation block.

    The previous OMML path rendered several constructs as broken glyphs in PDF.
    A borderless two-column table keeps the formula readable, wrapped, and numbered.
    """
    lines = FORMULA_LINES.get(num, [latex])
    tbl = cursor._parent.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(tbl)
    set_row_cant_split(tbl.rows[0])

    formula_cell, num_cell = tbl.rows[0].cells
    formula_cell.width = int(usable_emu * 0.88)
    num_cell.width = int(usable_emu * 0.12)
    num_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    first = formula_cell.paragraphs[0]
    first.paragraph_format.first_line_indent = 0
    first.paragraph_format.space_before = Pt(2)
    first.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        p = first if idx == 0 else formula_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(lines) == 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_formula_run(run)

    npara = num_cell.paragraphs[0]
    npara.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    npara.paragraph_format.first_line_indent = 0
    nrun = npara.add_run(f"({num})")
    set_formula_run(nrun)

    cursor._p.addnext(tbl._tbl)
    trail = OxmlElement("w:p")
    tbl._tbl.addnext(trail)
    return Paragraph(trail, cursor._parent)


def _set_cell_border(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    # 显式白底，覆盖任何 default 表格样式带来的灰底
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFFFF")
    tcPr.append(shd)
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tcPr.append(borders)


def insert_algo(doc, cursor: Paragraph, title: str, lines: list[str], usable_emu: int) -> Paragraph:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_row_cant_split(tbl.rows[0])
    cell = tbl.rows[0].cells[0]
    cell.width = usable_emu
    _set_cell_border(cell)
    # 关闭 tblPr 上 firstRow/lastRow/firstColumn 带来的"首行底纹"继承
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        for old in tblPr.findall(qn("w:tblLook")):
            tblPr.remove(old)
        look = OxmlElement("w:tblLook")
        for k, v in (("w:val", "0000"), ("w:firstRow", "0"), ("w:lastRow", "0"),
                     ("w:firstColumn", "0"), ("w:lastColumn", "0"),
                     ("w:noHBand", "0"), ("w:noVBand", "0")):
            look.set(qn(k), v)
        tblPr.append(look)
    tp = cell.paragraphs[0]
    tp.paragraph_format.space_after = Pt(2)
    trun = tp.add_run(title)
    trun.bold = True
    set_mono(trun, 10)
    for ln in lines:
        cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.first_line_indent = 0
        set_mono(cp.add_run(ln), 9.3)
    cursor._p.addnext(tbl._tbl)            # 把表移到锚点之后
    trail = OxmlElement("w:p")
    tbl._tbl.addnext(trail)
    return Paragraph(trail, cursor._parent)


def find_anchor(doc, substr: str) -> Paragraph:
    for p in doc.paragraphs:
        if substr in p.text:
            return p
    raise RuntimeError("anchor not found: " + substr)


def insert_image_block(doc, cursor: Paragraph, img: Paragraph, png: Path, width_inch: float) -> Paragraph:
    """在 cursor 之后插入（图片段 + 题注段），用 doc.add_picture 走默认图片 part。
    末尾返回题注段（作为下一段插入的锚点）。"""
    # 图片段
    p = insert_paragraph_after(cursor)
    p.alignment = 1  # CENTER
    r = p.add_run()
    r.add_picture(str(png), width=Inches(width_inch))
    return p


def insert_caption_after(doc, image_para: Paragraph, caption: str, ref: Paragraph) -> Paragraph:
    p = insert_paragraph_after(image_para)
    p.alignment = 1
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(8)
    pf.first_line_indent = 0
    run = p.add_run(caption)
    rr = _ref_run(ref)
    if rr is not None:
        copy_rpr(rr, run)
    run.bold = True
    return p


def replace_fig21(doc, png: Path) -> None:
    cap = next((p for p in doc.paragraphs if "事件知识图谱构建与时序推理总体框架" in p.text), None)
    assert cap is not None, "图2-1 题注未找到"
    cap_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text == cap.text)
    rid = None
    for j in range(cap_idx, -1, -1):
        blips = doc.paragraphs[j]._p.xpath(".//a:blip")
        if blips:
            rid = blips[0].get(qn("r:embed"))
            break
    assert rid, "图2-1 图片关系未找到"
    doc.part.related_parts[rid]._blob = png.read_bytes()
    b = png.read_bytes()
    w, h = struct.unpack(">II", b[16:24])
    for shp in doc.inline_shapes:
        bl = shp._inline.xpath(".//a:blip")
        if bl and bl[0].get(qn("r:embed")) == rid:
            target_w = Inches(5.9)
            shp.width = int(target_w)
            shp.height = int(target_w * h / w)
            break
    print(f"  替换图2-1 ← {png.name}  (png {w}x{h}, rid={rid})")


def reset_report_from_base() -> None:
    if BASE_REPORT.exists():
        shutil.copyfile(BASE_REPORT, REPORT)
        print(f"  还原基底 ← {BASE_REPORT.name}")


def bump_existing_figure_numbers(doc: Document) -> None:
    """Reserve 图2-2 for the inserted DNN-detail figure.

    The base report already has figures 图2-1..图2-10. The new detail figure is
    inserted after 图2-1, so all old references from 图2-2 onward move by +1.
    """
    replacements = {f"图2-{i}": f"__FIG2_{i + 1}__" for i in range(10, 1, -1)}
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            text = run.text
            for old, placeholder in replacements.items():
                text = text.replace(old, placeholder)
            for i in range(11, 2, -1):
                text = text.replace(f"__FIG2_{i}__", f"图2-{i}")
            run.text = text





def stabilize_tables(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            set_row_cant_split(row)
        if table.rows:
            set_row_repeat_header(table.rows[0])


# ----------------------------- 内容（公式/伪代码/价值段，按文档顺序） -----------------------------
EQ = lambda latex, num: ("eq", latex, num)        # noqa: E731
P = lambda text: ("p", text)                       # noqa: E731
ALGO = lambda title, lines: ("algo", title, lines)  # noqa: E731

SCENARIO = (
    "通俗地说，本课题面向这样一个实际场景：金融市场每天产生海量公告与新闻，人工逐篇研读既慢又"
    "易遗漏，而真正影响决策的，往往是其中“谁、在何时、对谁、做了什么”这类事件及其前后因果。"
    "本课题要做的，就是让机器自动把这些文本读成结构化、可核查的事件，连成一张可追溯的事件知识"
    "图谱，并据此对“接下来可能发生什么风险事件”给出带证据、带置信度、且错误率可控的预测，从而"
    "辅助舆情监测、风险预警与投资研究等决策。以一则股权质押公告为例：系统先抽出“某公司于某日"
    "质押若干股份”这一事件节点，再把它与同一主体的解除质押、股份减持等事件按时序与因果连成事件"
    "链，最后在该链上预测后续高风险事件，并给出带覆盖率保证的候选集合——其间每一步都保留可回溯"
    "到原文的证据，这正是本课题反复强调“可验证性”的现实意义所在。"
)

VALUE = (
    "从技术价值上看，本课题的贡献并不在于在某一环节堆叠更复杂的网络结构，而在于用“可验证性”这"
    "一统一原则把抽取、构图、推理三阶段连成一个可量化、可优化的整体，具体体现为三点：其一是可"
    "验证性奖励化，把格式合法、证据接地、全局一致与任务指标统一折算为强化学习奖励，区别于仅以"
    "任务 F1 或检索效用作为奖励的做法；其二是去价值网络的组相对策略与干预忠实度塑形，以组内相对"
    "优势替代额外的价值网络，并用反事实消融度量路径对预测的真实贡献，是路径强化学习在大模型时代"
    "的一次方法翻新；其三是漂移鲁棒的保形风险控制，在数据分布漂移下仍以有限样本保证维持目标覆盖"
    "率，为预测提供纯准确率方法普遍缺席的“错误率旋钮”。上述三点可统一表述为：在三阶段准确率目标"
    "之外，引入贯穿始终的可验证性约束，其总目标与验证器的三重身份分别形式化为式(2-1)与式(2-2)。"
)

# 每组 = (锚点子串, [块, ...])；块按列出顺序插在锚点之后。
GROUPS = [
    ("事件知识图谱构建与时序推理总体框架", [
        ("img", str(ARCH_DETAIL_PNG), 6.0, "图2-2  系统的详细技术框架（DNN 模块视角）"),
        P("该图从神经网络模块视角给出本系统的详细技术框架，把图2-1的五层数据流进一步展开为各阶段的算子层、张量形状与训练目标：抽取层由 QLoRA 微调的基座语言模型在 schema 约束下生成事件表，经模式检查与固定槽规范化输出事件节点；构图层对事件窗口采样 G 个候选输出、由复合可验证奖励 R(x,y) 评分、以组相对优势更新策略并由一致性求解器去除冲突边；推理层在时序图上以 MDP 形式采样路径，用反事实消融 F(ξ) 度量路径忠实度，并按可搜索权重与 RE-GCN 强基线融合排序；输出层用非一致性分数 a_i 在分位数 q_alpha 上构造预测集合 C(x)，并由 ACI 在线调整 alpha_{t+1} 把覆盖率维持在目标附近。验证器主线（顶端横向箭头）自上而下把门控/奖励/风险控制三类约束分别回传到对应模块，使每个阶段的输出都受到可验证约束的塑形。"),
    ]),

    ("它把文本中描述的事件抽取为节点", [P(SCENARIO)]),

    ("其二是诚实定位", [
        P(VALUE),
        EQ(r"\max_{M_1,M_2,M_3}\ \lambda_1 F1_{\mathrm{role}}+\lambda_2 F1_{\mathrm{rel}}+\lambda_3 \mathrm{MRR}-\lambda_4 \mathrm{Risk}\quad \mathrm{s.t.}\ g(v,D)=1,\ g(e,D)=1,\ \mathrm{Consistent}(E)=1,\ \Pr[y\in\mathcal{C}(x)]\ge 1-\alpha", "2-1"),
        EQ(r"\mathrm{Verifier}=\begin{cases}\mathrm{Gate}, & \text{extraction}\\ \mathrm{Reward}, & \text{relation}\\ \mathrm{RiskController}, & \text{inference}\end{cases}", "2-2"),
        P("式(2-1)的目标项对应三阶段的准确率指标，约束项对应可验证性：节点要有原文证据、边要有原文证据、图要全局一致、预测集合要满足覆盖率保证；式(2-2)表明同一验证器在抽取、关系、推理三阶段分别具体化为门控、奖励与风险控制器。"),
    ]),

    ("在实现上，SARGE以Qwen3-4B", [
        P("为形式化上述生成式抽取，下面给出 SARGE 的数学描述。一条事件记录定义为式(2-3)，其中 t 为事件类型、r 为角色、V_r 为取值、A_r 为取值在原文中的证据位置；文档级事件抽取被重述为受 schema 约束的事件表生成，即在合法事件表空间中求条件概率最大者，如式(2-4)；监督微调阶段以金标事件表为目标、仅对答案 token 计算交叉熵损失，如式(2-5)。"),
        EQ(r"e=\bigl(t,\ \{(r,V_r,A_r)\mid r\in\mathcal{R}_t\}\bigr)", "2-3"),
        EQ(r"\hat{Y}_D=\arg\max_{Y\in\mathcal{Y}(D,\mathcal{S})}\ p_\theta\bigl(Y\mid D,\mathcal{S},C_K(D),P_D\bigr)", "2-4"),
        EQ(r"\mathcal{L}_{\mathrm{SFT}}=-\sum_{(D,Y_D^\star)}\sum_{j=1}^{|Y_D^\star|}\log p_\theta\bigl(y_{D,j}^\star\mid x_D,\,y_{D,<j}^\star\bigr)", "2-5"),
        P("SARGE 从表面候选锚定到 schema 约束生成、再到模式检查与固定槽规范化的完整流程见算法2-1。"),
        ALGO("算法2-1  SARGE：schema 约束的事件表生成与规范化", [
            "输入：文档 D；事件模式 S=(类型, 角色集)；基座语言模型 π_θ",
            "输出：固定槽事件表 Ŷ_D",
            " 1: 构造有界表面候选 C_K(D)            // 金额/日期/机构/比例等按类型标注",
            " 2: 装配提示 x_D ← [D; S; C_K(D); 槽位先验 P_D; 格式指令]",
            " 3: 自回归生成事件表 Ŷ ← argmax_Y π_θ(Y | x_D)   // 训练时仅答案 token 计损",
            " 4: for 每条事件行 ê in Ŷ do",
            " 5:     if 事件类型/角色非法 或 字段为空 then 丢弃 ê   // Valid(ê, S)",
            " 6:     按基准模式排序角色、去除同一角色内重复取值",
            " 7: end for",
            " 8: 回填各论元的原文证据 span，输出规范化事件表 Ŷ_D",
        ]),
    ]),

    ("本节以字段级（角色级）micro平均F1为主指标", [
        P("上述字段级评测指标可形式化为式(2-6)，其中 TP、FP、FN 分别为预测正确、预测错误与漏报的角色字段数，仅当类型、角色名与取值均与金标完全一致时才计为 TP。"),
        EQ(r"P_{\mathrm{role}}=\frac{TP}{TP+FP},\qquad R_{\mathrm{role}}=\frac{TP}{TP+FN},\qquad F1_{\mathrm{role}}=\frac{2P_{\mathrm{role}}R_{\mathrm{role}}}{P_{\mathrm{role}}+R_{\mathrm{role}}}", "2-6"),
    ]),

    ("为刻画角色取值恢复与一致事件行组装之间的差距", [
        P("为量化角色级成功与整行成功之间的差距，定义精确记录匹配率 ExactRec 及记录绑定差 Δbind，如式(2-7)，其中 M_{D,t} 为同文档同类型下预测记录与金标记录的一对一匹配。Δbind 越大，说明模型越容易把多条同类事件的论元组装进错误的事件行，即同质记录绑定问题越突出。"),
        EQ(r"\mathrm{ExactRec}=\frac{2\sum_{D,t}\sum_{(e,\hat{e})\in\mathcal{M}_{D,t}}\mathbf{1}[e=\hat{e}]}{\sum_{D,t}\bigl(|Y_{D,t}|+|\hat{Y}_{D,t}|\bigr)},\qquad \Delta_{\mathrm{bind}}=F1_{\mathrm{role}}-\mathrm{ExactRec}", "2-7"),
    ]),

    ("具体而言，方法分为监督微调与强化优化两个阶段", [
        P("形式上，对输入事件窗口 x 与模型输出 y，复合可验证奖励定义为式(2-8)，其格式、证据接地、全局一致性与任务 F1 四个分量分别如式(2-9)所示；GRPO 对同一输入采样 G 个候选，以组内标准化优势式(2-10)替代价值网络，并以带裁剪与 KL 锚定的策略目标式(2-11)进行优化，其中 ρ_i=π_θ(y_i|x)/π_{θ_old}(y_i|x) 为重要性比、参考策略 π_ref 取监督微调模型。完整训练流程见算法2-2。"),
        EQ(r"R(x,y)=w_f R_{\mathrm{format}}+w_g R_{\mathrm{ground}}+w_c R_{\mathrm{consist}}+w_t R_{\mathrm{task}}", "2-8"),
        EQ(r"\begin{aligned}R_{\mathrm{format}}&=\mathbf{1}[\mathrm{Parse}(y)\neq\varnothing]\cdot\frac{|\mathrm{ValidEdges}(y)|}{\max(|\mathrm{Edges}(y)|,1)},\\ R_{\mathrm{ground}}&=\frac{1}{\max(|\hat{E}|,1)}\sum_{e\in\hat{E}}\mathbf{1}[\eta_e\subseteq D],\\ R_{\mathrm{consist}}&=1-\frac{|\hat{E}\setminus\Pi(\hat{E})|}{\max(|\hat{E}|,1)},\qquad R_{\mathrm{task}}=F1_{\mathrm{rel}}(\hat{E},E^\star)\end{aligned}", "2-9"),
        EQ(r"\hat{A}_i=\frac{R_i-\mathrm{mean}(R_1,\dots,R_G)}{\max(\mathrm{std}(R_1,\dots,R_G),\,\epsilon)}", "2-10"),
        EQ(r"\mathcal{J}_{\mathrm{GRPO}}(\theta)=\frac{1}{G}\sum_{i=1}^{G}\min\bigl(\rho_i\hat{A}_i,\ \mathrm{clip}(\rho_i,1-\epsilon_c,1+\epsilon_c)\hat{A}_i\bigr)-\beta\,D_{\mathrm{KL}}(\pi_\theta\|\pi_{\mathrm{ref}})", "2-11"),
        ALGO("算法2-2  GRPO-RLVR：可验证奖励的事件关系抽取训练", [
            "输入：事件窗口 x；SFT 参考策略 π_ref；组大小 G；KL 系数 β",
            "输出：关系抽取策略 π_θ",
            " 1: π_θ ← π_ref                          // 监督微调初始化",
            " 2: for 每个训练步 do",
            " 3:     采样一组候选输出 {y_1,…,y_G} ~ π_θ(· | x)",
            " 4:     for i = 1..G do",
            " 5:         R_i ← w_f·R_format + w_g·R_ground + w_c·R_consist + w_t·R_task",
            " 6:     end for",
            " 7:     Â_i ← (R_i − mean(R)) / max(std(R), ε)    // 组相对优势, 无 critic",
            " 8:     θ ← argmax Σ_i min(ρ_i Â_i, clip(ρ_i,1±ε_c) Â_i) − β·KL(π_θ‖π_ref)",
            " 9: end for",
        ]),
    ]),

    ("在得到四类关系的抽取结果后，还需经一致性求解把局部判断整合", [
        P("一致性约束可形式化如下：共指被视为等价关系、时序“先于”满足传递性，并对时序边取传递闭包，如式(2-12)；此外，在图谱构建阶段即可施加保形风险控制做边准入——以漏报率 FNR 为风险，按式(2-13)选取满足有限样本上界的最紧阈值 τ*，从而在召回风险受控的前提下过滤低置信边、提高图谱精度。"),
        EQ(r"\begin{aligned}&v_i\sim v_j\Rightarrow v_j\sim v_i,\quad (v_i\sim v_j\wedge v_j\sim v_k)\Rightarrow v_i\sim v_k,\\ &(v_i\prec v_j\wedge v_j\prec v_k)\Rightarrow v_i\prec v_k,\quad E_{\mathrm{temp}}^{+}=\mathrm{TransitiveClosure}(E_{\mathrm{temp}})\end{aligned}", "2-12"),
        EQ(r"\widehat{\mathrm{FNR}}(\tau)=\frac{1}{n}\sum_{i=1}^{n}\mathbf{1}[s_{e_i^\star}<\tau],\quad U_{\mathrm{CRC}}(\tau)=\frac{n}{n+1}\widehat{\mathrm{FNR}}(\tau)+\frac{1}{n+1},\quad \tau^\star=\max\{\tau:U_{\mathrm{CRC}}(\tau)\le\alpha\}", "2-13"),
    ]),

    ("在方法上，本节把时序预测建模为图上的路径搜索", [
        P("形式化地，给定历史四元组集合与未来查询，预测模型为每个候选实体计算式(2-14)的得分；路径推理把预测建模为时序图上的马尔可夫决策过程，以查询主体为起点，状态与动作如式(2-15)所示；轨迹奖励综合终点命中、路径忠实度与势函数塑形，如式(2-16)，其中忠实度 F(ξ) 通过反事实消融度量，如式(2-17)：若移除路径后正确答案得分明显下降，则该路径对预测有实质贡献。路径采样与训练流程见算法2-3。"),
        EQ(r"f_\theta\bigl(o\mid q,\mathcal{H}_{<\tau_q}\bigr),\qquad \mathcal{H}_{<\tau_q}=\{(s,r,o,\tau)\in\mathcal{H}\mid \tau<\tau_q\}", "2-14"),
        EQ(r"\begin{aligned}s_l&=(q,\,e_l,\,\mathrm{path}_{0:l}),\qquad e_0=s_q,\\ \mathcal{A}(s_l)&=\{(r,e',\tau'):(e_l,r,e',\tau')\in\mathcal{H}_{<\tau_q}\}\cup\{\mathrm{STAY}\}\end{aligned}", "2-15"),
        EQ(r"R(\xi)=c_h\mathbf{1}[e_L=o^\star]+c_f F(\xi)\,\mathbf{1}[e_L=o^\star]+c_s\sum_{l=0}^{L-1}\bigl(\gamma\Phi(s_{l+1})-\Phi(s_l)\bigr)", "2-16"),
        EQ(r"F(\xi)=\max\bigl(0,\ f_\theta(o^\star\mid q,\mathcal{H})-f_\theta(o^\star\mid q,\mathcal{H}\setminus\xi)\bigr)", "2-17"),
        ALGO("算法2-3  Path-RL：可回溯路径推理与忠实度塑形", [
            "输入：查询 q=(s_q, r_q, ?, τ_q)；历史图 H_{<τq}；步长 T；组大小 G",
            "输出：候选实体排序与证据路径",
            " 1: for g = 1..G do",
            " 2:     e_0 ← s_q;  path ← ∅",
            " 3:     for l = 0..T−1 do",
            " 4:         A(s_l) ← {(r,e',τ'):(e_l,r,e',τ')∈H_{<τq}, 按近因取 top-K} ∪ {STAY}",
            " 5:         采样动作 a_l ~ π_θ(· | s_l), 转移到 e_{l+1}",
            " 6:     end for",
            " 7:     R(ξ) ← c_h·1[e_T=o*] + c_f·F(ξ)·1[e_T=o*] + c_s·Σ_l(γΦ(s_{l+1})−Φ(s_l))",
            " 8:           // F(ξ) = 消融该路径后正确答案得分的下降幅度（干预忠实度）",
            " 9: end for",
            "10: 组内标准化优势 + REINFORCE 更新 θ；聚合各 rollout 终点得分得到排序",
        ]),
    ]),

    ("在评价指标上，本节采用时序知识图谱推理常用的MRR", [
        P("上述排序质量以平均倒数排名 MRR 与 Hits@k 衡量，如式(2-18)；二者均在时间感知过滤设定下计算，即排名时屏蔽历史中已知为真的其他四元组，如式(2-19)，以避免把已知事实误判为错误负例。"),
        EQ(r"\mathrm{MRR}=\frac{1}{N}\sum_{i=1}^{N}\frac{1}{\mathrm{rank}_i},\qquad \mathrm{Hits@}k=\frac{1}{N}\sum_{i=1}^{N}\mathbf{1}[\mathrm{rank}_i\le k]", "2-18"),
        EQ(r"\mathrm{rank}^{\mathrm{tfilt}}(o^\star)=1+\sum_{o\in\mathcal{V}\setminus\mathcal{F}(q)}\mathbf{1}\bigl[f_\theta(o\mid q)>f_\theta(o^\star\mid q)\bigr]", "2-19"),
    ]),

    ("那么风险控制才是本节试图提供", [
        P("形式化地，保形预测把单点排序转化为带覆盖保证的预测集合：以正确答案排名为非一致性分数，预测集合与静态 split 校准阈值如式(2-20)；为应对分布漂移，自适应保形按式(2-21)用在线误覆盖反馈调整错误率参数 α_t；实际覆盖率、平均集合大小与漂移覆盖失真 DriftGap 的度量如式(2-22)，DriftGap 越小说明各时间段覆盖率越贴近目标。四种校准策略（split/ACI/weighted/CRC）的统一在线流程见算法2-4。"),
        EQ(r"\mathcal{C}_{q_\alpha}(x)=\{o\in\mathcal{V}:\mathrm{rank}(o\mid x)\le q_\alpha\},\qquad q_\alpha=a_{(\lceil(n+1)(1-\alpha)\rceil)}", "2-20"),
        EQ(r"\alpha_{t+1}=\alpha_t+\gamma(\alpha-\mathrm{err}_t),\qquad \mathrm{err}_t=\mathbf{1}[y_t\notin\mathcal{C}_t(x_t)]", "2-21"),
        EQ(r"\widehat{\mathrm{Cov}}=\frac{1}{N}\sum_{t=1}^{N}\mathbf{1}[y_t\in\mathcal{C}_t(x_t)],\quad \widehat{\mathrm{Size}}=\frac{1}{N}\sum_{t=1}^{N}|\mathcal{C}_t(x_t)|,\quad \mathrm{DriftGap}=\max_{W\in\mathcal{W}}\left|\tfrac{1}{|W|}\sum_{t\in W}\mathbf{1}[y_t\in\mathcal{C}_t(x_t)]-(1-\alpha)\right|", "2-22"),
        ALGO("算法2-4  自适应保形风险控制（split / ACI / weighted / CRC）", [
            "输入：在线校准分数流（非一致性 = gold 排名）；目标错误率 α；步长 γ",
            "输出：每步预测集合 C_t 与在线覆盖保证",
            " 1: 初始化阈值 q ← 校准集第 ⌈(n+1)(1−α)⌉ 小的分数",
            " 2: for 每个到达的查询 t do",
            " 3:     C_t(x_t) ← { o : rank(o | x_t) ≤ q }        // 预测集合",
            " 4:     揭示真值 y_t,  err_t ← 1[y_t ∉ C_t]",
            " 5:     ACI:       α_t ← α_t + γ(α − err_t),  q ← (1−α_t) 分位",
            " 6:     weighted:  q ← 近因加权 (1−α) 分位（权重几何衰减）",
            " 7:     CRC:       q ← min{ k : (n/(n+1))·R̂_n(k) + B/(n+1) ≤ α }",
            " 8: end for",
        ]),
    ]),
]


def main() -> None:
    print("打开:", REPORT.name)
    reset_report_from_base()
    doc = Document(str(REPORT))
    bump_existing_figure_numbers(doc)
    sec = doc.sections[0]
    usable = int(sec.page_width - sec.left_margin - sec.right_margin)

    n_eq = n_algo = n_p = 0
    for anchor_sub, blocks in GROUPS:
        ref = find_anchor(doc, anchor_sub)
        cursor = ref
        for blk in blocks:
            kind = blk[0]
            if kind == "p":
                cursor = insert_body(cursor, blk[1], ref)
                n_p += 1
            elif kind == "eq":
                cursor = insert_equation(cursor, blk[1], blk[2], usable)
                n_eq += 1
            elif kind == "algo":
                cursor = insert_algo(doc, cursor, blk[1], blk[2], usable)
                n_algo += 1
            elif kind == "img":
                _, png_path, width_in, caption = blk
                image_para = insert_image_block(doc, cursor, ref, Path(png_path), width_in)
                cursor = insert_caption_after(doc, image_para, caption, ref)
        print(f"  ✓ 锚点[{anchor_sub[:16]}…] 插入 {len(blocks)} 块")

    print("替换图2-1 为分层架构图…")
    replace_fig21(doc, ARCH_PNG)
    stabilize_tables(doc)

    doc.save(str(REPORT))
    print(f"已保存。新增 段落 {n_p} · 公式 {n_eq} · 算法 {n_algo}")


if __name__ == "__main__":
    main()
